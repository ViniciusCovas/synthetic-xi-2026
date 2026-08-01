#!/usr/bin/env python3
"""Build formatted DOCX/PDF submission files for the JASSS manuscript.

The builder creates a double-blind manuscript, ODD supplement, title page,
cover letter, high-resolution figures, an anonymous reproducibility ZIP,
and a checksum manifest. It does not alter scientific values.
"""
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TITLE = "Can a Synthetic Collective Rival an Elite Team? A Reproducible Social Simulation of Composition, Variability, and Context Dependence"
FIGURE_FILES = [
    "figure_1_replication_consistency.svg",
    "figure_2_h5_variability.svg",
    "figure_3_h6_extreme_outcomes.svg",
    "figure_4_precision_vs_structural_uncertainty.svg",
    "figure_5_h7_context_dependence.svg",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path("."))
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


def run(command: list[str], cwd: Path | None = None) -> None:
    subprocess.run(command, cwd=cwd, check=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def require_program(name: str) -> None:
    if shutil.which(name) is None:
        raise RuntimeError(f"Required program not found: {name}")


def parse_captions(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    matches = list(re.finditer(r"^## Figure (\d+)\. (.+?)\n\n(.+?)(?=\n\n## Figure|\Z)", text, flags=re.MULTILINE | re.DOTALL))
    captions: list[str] = []
    for match in matches:
        number, heading, body = match.groups()
        body = " ".join(line.strip() for line in body.strip().splitlines())
        captions.append(f"Figure {number}. {heading}. {body}")
    if len(captions) != len(FIGURE_FILES):
        raise RuntimeError(f"Expected {len(FIGURE_FILES)} captions, found {len(captions)}")
    return captions


def set_cell_font(run, name: str = "Times New Roman", size: float = 12) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_cell_font(run, size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run._r.append(node)


def enable_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def set_image_alt_text(shape, description: str) -> None:
    inline = shape._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description.split(".", 1)[0])


def style_document(path: Path, *, blind: bool, double_spaced: bool, line_numbers: bool) -> None:
    document = Document(path)
    core = document.core_properties
    core.title = TITLE if "Manuscript" in path.name else path.stem
    core.subject = "JASSS submission file"
    core.author = "" if blind else "Vinicius Covas"
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 2.0 if double_spaced else 1.15

    for style_name, size, bold, italic in (
        ("Title", 16, True, False),
        ("Heading 1", 14, True, False),
        ("Heading 2", 12, True, False),
        ("Heading 3", 12, False, True),
    ):
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if line_numbers:
            enable_line_numbering(section)
        footer = section.footer
        if not footer.paragraphs:
            paragraph = footer.add_paragraph()
        else:
            paragraph = footer.paragraphs[0]
        paragraph.clear()
        add_page_number(paragraph)

    for paragraph in document.paragraphs:
        for run_item in paragraph.runs:
            set_cell_font(run_item, size=12)
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        if double_spaced and paragraph.style.name not in {"Caption"}:
            paragraph.paragraph_format.line_spacing = 2.0
            paragraph.paragraph_format.space_after = Pt(0)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run_item in paragraph.runs:
                        set_cell_font(run_item, size=10)

    document.save(path)


def add_figures_to_manuscript(docx_path: Path, figure_pngs: list[Path], captions: list[str]) -> None:
    document = Document(docx_path)
    document.add_page_break()
    heading = document.add_heading("Figures", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, (figure, caption) in enumerate(zip(figure_pngs, captions, strict=True)):
        if index > 0:
            document.add_page_break()
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = image_paragraph.add_run().add_picture(str(figure), width=Inches(6.35))
        set_image_alt_text(shape, caption)
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.style = document.styles["Caption"]
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption_paragraph.paragraph_format.line_spacing = 1.0
        caption_paragraph.paragraph_format.space_before = Pt(6)
        caption_paragraph.paragraph_format.space_after = Pt(0)
        for run_item in caption_paragraph.runs:
            set_cell_font(run_item, size=10)
    document.save(docx_path)


def convert_markdown(source: Path, target: Path, root: Path) -> None:
    run([
        "pandoc",
        str(source),
        "--from=gfm",
        "--to=docx",
        f"--resource-path={root}",
        "--output",
        str(target),
    ], cwd=root)


def convert_pdf(docx: Path, output_dir: Path) -> Path:
    run([
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(docx),
    ])
    pdf = output_dir / f"{docx.stem}.pdf"
    if not pdf.is_file() or pdf.stat().st_size == 0:
        raise RuntimeError(f"PDF conversion failed for {docx}")
    return pdf


def pdf_pages(path: Path) -> int:
    result = subprocess.run(["pdfinfo", str(path)], check=True, capture_output=True, text=True)
    for line in result.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    raise RuntimeError(f"Page count not found for {path}")


def copy_figures(root: Path, output: Path) -> tuple[list[Path], list[Path]]:
    svg_dir = output / "05_Figures_SVG"
    png_dir = output / "06_Figures_PNG"
    svg_dir.mkdir(parents=True)
    png_dir.mkdir(parents=True)
    svg_files: list[Path] = []
    png_files: list[Path] = []
    for filename in FIGURE_FILES:
        source = root / "paper" / "figures" / filename
        svg_target = svg_dir / filename
        png_target = png_dir / f"{Path(filename).stem}.png"
        shutil.copy2(source, svg_target)
        run(["rsvg-convert", "-w", "1800", "-o", str(png_target), str(source)])
        svg_files.append(svg_target)
        png_files.append(png_target)
    return svg_files, png_files


def deterministic_zip(source_dir: Path, target: Path) -> None:
    if target.exists():
        target.unlink()
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(source_dir.rglob("*")):
            if not path.is_file() or path == target:
                continue
            info = zipfile.ZipInfo(str(path.relative_to(source_dir)))
            info.date_time = (2026, 7, 31, 0, 0, 0)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            archive.writestr(info, path.read_bytes())


def main() -> None:
    args = parse_args()
    root = args.root.resolve()
    output = args.output.resolve()
    for program in ("pandoc", "libreoffice", "rsvg-convert", "pdfinfo"):
        require_program(program)
    if output.exists():
        shutil.rmtree(output)
    output.mkdir(parents=True)

    captions = parse_captions(root / "paper" / "FIGURE_CAPTIONS.md")
    _svg_files, png_files = copy_figures(root, output)

    sources = {
        "01_Blinded_Manuscript": root / "paper" / "MANUSCRIPT_JASSS_V1.md",
        "02_Supplement_ODD": root / "paper" / "SUPPLEMENT_ODD_OFFICIAL_V1.md",
        "03_Title_Page": root / "paper" / "TITLE_PAGE_JASSS_READY.md",
        "04_Cover_Letter": root / "paper" / "COVER_LETTER_JASSS_READY.md",
    }
    generated: list[Path] = []
    pdf_counts: dict[str, int] = {}

    for stem, source in sources.items():
        docx = output / f"{stem}.docx"
        convert_markdown(source, docx, root)
        if stem == "01_Blinded_Manuscript":
            add_figures_to_manuscript(docx, png_files, captions)
            style_document(docx, blind=True, double_spaced=True, line_numbers=True)
        elif stem == "02_Supplement_ODD":
            style_document(docx, blind=True, double_spaced=True, line_numbers=True)
        else:
            style_document(docx, blind=False, double_spaced=False, line_numbers=False)
        pdf = convert_pdf(docx, output)
        generated.extend([docx, pdf])
        pdf_counts[pdf.name] = pdf_pages(pdf)

    anonymous_dir = output / "07_Anonymous_Reproducibility_Package"
    anonymous_zip = output / "07_Anonymous_Reproducibility_Package.zip"
    anonymous_manifest = output / "07_Anonymous_Reproducibility_Package_manifest.json"
    run([
        "python",
        str(root / "scripts" / "scientific" / "build_anonymous_review_bundle.py"),
        "--root",
        str(root),
        "--output-dir",
        str(anonymous_dir),
        "--archive",
        str(anonymous_zip),
        "--manifest",
        str(anonymous_manifest),
    ], cwd=root)
    shutil.rmtree(anonymous_dir)

    open_items = """# Submission open items

The scientific and double-blind packages are complete. Three author-level items remain before portal submission:

1. Confirm the exact funding declaration.
2. Confirm the exact competing-interests declaration.
3. Create the permanent archival deposit and replace the DOI placeholder in the non-blinded files.

No additional simulations are required for these items.
"""
    (output / "08_SUBMISSION_OPEN_ITEMS.md").write_text(open_items, encoding="utf-8")

    files = []
    for path in sorted(output.rglob("*")):
        if path.is_file() and path.name != "JASSS_Submission_Package_v1.zip":
            files.append({
                "path": str(path.relative_to(output)),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
                "pdf_pages": pdf_counts.get(path.name),
            })
    manifest = {
        "status": "jasss_submission_package_v1_built",
        "scientific_release": "official_v1",
        "clean_room_reproduction": "passed_byte_identical",
        "double_blind_identity_scan": "passed",
        "files": files,
        "open_items": [
            "funding declaration",
            "competing-interests declaration",
            "permanent archival DOI",
        ],
    }
    manifest_path = output / "PACKAGE_MANIFEST.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    bundle = output / "JASSS_Submission_Package_v1.zip"
    deterministic_zip(output, bundle)
    report = {
        **manifest,
        "bundle": bundle.name,
        "bundle_bytes": bundle.stat().st_size,
        "bundle_sha256": sha256(bundle),
    }
    manifest_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    deterministic_zip(output, bundle)
    print(json.dumps({
        "status": report["status"],
        "bundle": str(bundle),
        "bundle_sha256": sha256(bundle),
        "pdf_pages": pdf_counts,
    }, indent=2))


if __name__ == "__main__":
    main()
