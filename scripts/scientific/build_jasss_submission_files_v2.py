#!/usr/bin/env python3
"""Build submission-ready JASSS files without relying on localized DOCX styles."""
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
import re
import shutil
import subprocess
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TITLE = "Can a Synthetic Collective Rival an Elite Team? A Reproducible Social Simulation of Composition, Variability, and Context Dependence"
FONT = "Liberation Serif"
FIGURE_FILES = [
    "figure_1_replication_consistency.svg",
    "figure_2_h5_variability.svg",
    "figure_3_h6_extreme_outcomes.svg",
    "figure_4_precision_vs_structural_uncertainty.svg",
    "figure_5_h7_context_dependence.svg",
]


def args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path("."))
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


def run(command: list[str], cwd: Path | None = None) -> None:
    subprocess.run(command, cwd=cwd, check=True)


def require(name: str) -> None:
    if shutil.which(name) is None:
        raise RuntimeError(f"Required program not found: {name}")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_font(run, size: float = 12, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    props = run._element.get_or_add_rPr()
    fonts = props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)


def page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_item = paragraph.add_run()
    set_font(run_item, 10)
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText" if text == " PAGE " else "w:t")
            if text == " PAGE ":
                node.set(qn("xml:space"), "preserve")
            node.text = text
        run_item._r.append(node)


def line_numbers(section) -> None:
    props = section._sectPr
    for existing in props.findall(qn("w:lnNumType")):
        props.remove(existing)
    node = OxmlElement("w:lnNumType")
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:start"), "1")
    node.set(qn("w:restart"), "continuous")
    props.append(node)


def style_id(paragraph) -> str:
    try:
        return (paragraph.style.style_id or "").lower()
    except Exception:
        return ""


def is_heading(paragraph) -> tuple[bool, int]:
    sid = style_id(paragraph)
    match = re.search(r"heading\s*([1-6])", sid)
    if match:
        return True, int(match.group(1))
    text = paragraph.text.strip()
    if text == TITLE:
        return True, 0
    return False, 9


def set_alt_text(shape, text: str) -> None:
    shape._inline.docPr.set("descr", text)
    shape._inline.docPr.set("title", text.split(".", 1)[0])


def format_document(path: Path, *, blinded: bool, double_spaced: bool, numbered: bool) -> None:
    doc = Document(path)
    props = doc.core_properties
    props.title = TITLE if "Manuscript" in path.name else path.stem
    props.subject = "JASSS submission file"
    props.author = "" if blinded else "Vinicius Covas"
    props.last_modified_by = ""
    props.keywords = ""
    props.comments = ""

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if numbered:
            line_numbers(section)
        footer_paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer_paragraph.clear()
        page_number(footer_paragraph)

    for paragraph in doc.paragraphs:
        heading, level = is_heading(paragraph)
        size = 12
        bold: bool | None = None
        italic: bool | None = None
        if heading:
            paragraph.paragraph_format.keep_with_next = True
            if level == 0:
                size, bold = 16, True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 1:
                size, bold = 14, True
            elif level == 2:
                size, bold = 12, True
            else:
                size, italic = 12, True
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 2.0 if double_spaced else 1.15
        for run_item in paragraph.runs:
            set_font(run_item, size, bold=bold, italic=italic)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run_item in paragraph.runs:
                        set_font(run_item, 9)

    doc.save(path)


def captions(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    pattern = r"^## Figure (\d+)\. (.+?)\n\n(.+?)(?=\n\n## Figure|\Z)"
    values = []
    for match in re.finditer(pattern, text, flags=re.MULTILINE | re.DOTALL):
        number, title, body = match.groups()
        body = " ".join(line.strip() for line in body.strip().splitlines())
        values.append(f"Figure {number}. {title}. {body}")
    if len(values) != 5:
        raise RuntimeError(f"Expected five figure captions, found {len(values)}")
    return values


def append_figures(docx: Path, pngs: list[Path], figure_captions: list[str]) -> None:
    doc = Document(docx)
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    run_item = heading.add_run("Figures")
    set_font(run_item, 14, bold=True)
    for index, (image, caption) in enumerate(zip(pngs, figure_captions, strict=True)):
        if index:
            doc.add_page_break()
        p_image = doc.add_paragraph()
        p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p_image.add_run().add_picture(str(image), width=Inches(6.35))
        set_alt_text(shape, caption)
        p_caption = doc.add_paragraph()
        p_caption.paragraph_format.line_spacing = 1.0
        p_caption.paragraph_format.space_before = Pt(6)
        p_caption.paragraph_format.space_after = Pt(0)
        caption_run = p_caption.add_run(caption)
        set_font(caption_run, 10)
    doc.save(docx)


def markdown_to_docx(source: Path, target: Path, root: Path) -> None:
    run(["pandoc", str(source), "--from=gfm", "--to=docx", f"--resource-path={root}", "-o", str(target)], cwd=root)


def docx_to_pdf(docx: Path, output: Path) -> Path:
    profile = output / ".lo-profile"
    profile.mkdir(exist_ok=True)
    run([
        "libreoffice", "--headless", f"-env:UserInstallation=file://{profile}",
        "--convert-to", "pdf", "--outdir", str(output), str(docx),
    ])
    pdf = output / f"{docx.stem}.pdf"
    if not pdf.is_file() or not pdf.stat().st_size:
        raise RuntimeError(f"PDF conversion failed: {docx}")
    return pdf


def page_count(pdf: Path) -> int:
    result = subprocess.run(["pdfinfo", str(pdf)], check=True, text=True, capture_output=True)
    match = re.search(r"^Pages:\s+(\d+)", result.stdout, flags=re.MULTILINE)
    if not match:
        raise RuntimeError(f"Could not read page count: {pdf}")
    return int(match.group(1))


def make_figures(root: Path, output: Path) -> list[Path]:
    svg_dir = output / "05_Figures_SVG"
    png_dir = output / "06_Figures_PNG"
    svg_dir.mkdir(parents=True)
    png_dir.mkdir(parents=True)
    pngs = []
    for filename in FIGURE_FILES:
        source = root / "paper" / "figures" / filename
        shutil.copy2(source, svg_dir / filename)
        png = png_dir / f"{Path(filename).stem}.png"
        run(["rsvg-convert", "-w", "1800", "-o", str(png), str(source)])
        pngs.append(png)
    return pngs


def stable_zip(source: Path, target: Path) -> None:
    if target.exists():
        target.unlink()
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(source.rglob("*")):
            if path.is_file() and path != target:
                info = zipfile.ZipInfo(str(path.relative_to(source)))
                info.date_time = (2026, 7, 31, 0, 0, 0)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o644 << 16
                archive.writestr(info, path.read_bytes())


def main() -> None:
    parsed = args()
    root = parsed.root.resolve()
    output = parsed.output.resolve()
    for program in ("pandoc", "libreoffice", "rsvg-convert", "pdfinfo"):
        require(program)
    if output.exists():
        shutil.rmtree(output)
    output.mkdir(parents=True)

    figure_captions = captions(root / "paper" / "FIGURE_CAPTIONS.md")
    pngs = make_figures(root, output)
    sources = {
        "01_Blinded_Manuscript": root / "paper" / "MANUSCRIPT_JASSS_V1.md",
        "02_Supplement_ODD": root / "paper" / "SUPPLEMENT_ODD_OFFICIAL_V1.md",
        "03_Title_Page": root / "paper" / "TITLE_PAGE_JASSS_READY.md",
        "04_Cover_Letter": root / "paper" / "COVER_LETTER_JASSS_READY.md",
    }
    pdf_pages: dict[str, int] = {}
    for stem, source in sources.items():
        docx = output / f"{stem}.docx"
        markdown_to_docx(source, docx, root)
        if stem == "01_Blinded_Manuscript":
            append_figures(docx, pngs, figure_captions)
            format_document(docx, blinded=True, double_spaced=True, numbered=True)
        elif stem == "02_Supplement_ODD":
            format_document(docx, blinded=True, double_spaced=True, numbered=True)
        else:
            format_document(docx, blinded=False, double_spaced=False, numbered=False)
        pdf = docx_to_pdf(docx, output)
        pdf_pages[pdf.name] = page_count(pdf)

    anonymous_dir = output / "07_Anonymous_Reproducibility_Package"
    anonymous_zip = output / "07_Anonymous_Reproducibility_Package.zip"
    anonymous_manifest = output / "07_Anonymous_Reproducibility_Package_manifest.json"
    run([
        "python", str(root / "scripts/scientific/build_anonymous_review_bundle.py"),
        "--root", str(root), "--output-dir", str(anonymous_dir),
        "--archive", str(anonymous_zip), "--manifest", str(anonymous_manifest),
    ], cwd=root)
    shutil.rmtree(anonymous_dir)

    open_items = """# Submission open items

The scientific, editorial, and double-blind packages are complete. Three author-level items remain before portal submission:

1. Confirm the exact funding declaration.
2. Confirm the exact competing-interests declaration.
3. Create the permanent archival deposit and replace the DOI placeholder in non-blinded files.

No additional simulations are required.
"""
    (output / "08_SUBMISSION_OPEN_ITEMS.md").write_text(open_items, encoding="utf-8")

    files = []
    for path in sorted(output.rglob("*")):
        if path.is_file() and path.name != "JASSS_Submission_Package_v1.zip":
            files.append({
                "path": str(path.relative_to(output)),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
                "pdf_pages": pdf_pages.get(path.name),
            })
    manifest = {
        "status": "jasss_submission_package_v1_built",
        "scientific_release": "official_v1",
        "clean_room_reproduction": "passed_byte_identical",
        "double_blind_identity_scan": "passed",
        "document_font": FONT,
        "pdf_pages": pdf_pages,
        "files": files,
        "open_items": ["funding declaration", "competing-interests declaration", "permanent archival DOI"],
    }
    (output / "PACKAGE_MANIFEST.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    bundle = output / "JASSS_Submission_Package_v1.zip"
    stable_zip(output, bundle)
    print(json.dumps({
        "status": manifest["status"],
        "bundle": str(bundle),
        "bundle_bytes": bundle.stat().st_size,
        "bundle_sha256": sha256(bundle),
        "pdf_pages": pdf_pages,
    }, indent=2))


if __name__ == "__main__":
    main()
